"""Parse a CV/resume (PDF, DOCX, or plain text) into raw text and a structured profile."""
from __future__ import annotations

import re
from dataclasses import dataclass, field, asdict
from io import BytesIO
from pathlib import Path
from typing import BinaryIO

from .skills import extract_skills


@dataclass
class CVProfile:
    raw_text: str = ""
    name: str | None = None                               # candidate name (for letter signatures)
    skills: list[str] = field(default_factory=list)
    titles: list[str] = field(default_factory=list)      # likely job titles found
    years_experience: int | None = None
    location: str | None = None
    seniority: str | None = None                          # junior / mid / senior / lead
    suggested_keywords: str = ""                          # search query seed for job sources

    def to_dict(self) -> dict:
        return asdict(self)


# ---------------------------------------------------------------------------
# Text extraction
# ---------------------------------------------------------------------------

def extract_text(path: str | Path) -> str:
    """Extract plain text from a PDF, DOCX, or text file."""
    path = Path(path)
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return _extract_pdf(path)
    if suffix in (".docx",):
        return _extract_docx(path)
    if suffix in (".txt", ".md", ".text", ""):
        return path.read_text(encoding="utf-8", errors="ignore")
    if suffix == ".doc":
        raise ValueError("Legacy .doc files are not supported — please save as .docx or PDF.")
    # Best-effort fallback: try reading as text.
    return path.read_text(encoding="utf-8", errors="ignore")


def extract_text_from_bytes(data: bytes, filename: str) -> str:
    """Extract text from uploaded bytes — fully in memory, never touching disk.

    (Earlier this spooled to a NamedTemporaryFile(delete=False); a crash mid-parse
    could leave a plaintext CV in the OS temp dir — a privacy leak. Parsing from a
    BytesIO keeps the CV in process memory only.)
    """
    suffix = Path(filename).suffix.lower()
    if suffix == ".pdf":
        return _extract_pdf_stream(BytesIO(data), name=filename)
    if suffix == ".docx":
        return _extract_docx_stream(BytesIO(data))
    if suffix == ".doc":
        raise ValueError("Legacy .doc files are not supported — please save as .docx or PDF.")
    return data.decode("utf-8", errors="ignore")


def _extract_pdf(path: Path) -> str:
    with open(path, "rb") as f:
        return _extract_pdf_stream(f, name=path.name)


def _extract_pdf_stream(stream: BinaryIO, name: str = "document.pdf") -> str:
    # Prefer pypdf (pure-python, light); fall back to pdfplumber if available.
    first_error: Exception | None = None
    try:
        from pypdf import PdfReader
        reader = PdfReader(stream)
        return "\n".join((page.extract_text() or "") for page in reader.pages)
    except Exception as e:
        first_error = e  # remember the real cause so we don't misattribute it below
    try:
        import pdfplumber
        stream.seek(0)
        with pdfplumber.open(stream) as pdf:
            return "\n".join((page.extract_text() or "") for page in pdf.pages)
    except Exception as e2:
        raise RuntimeError(
            f"Could not read PDF '{name}'. pypdf failed ({first_error}); "
            f"pdfplumber failed ({e2}). If the file is a scanned image, paste your CV text instead."
        )


def _extract_docx(path: Path) -> str:
    with open(path, "rb") as f:
        return _extract_docx_stream(f)


def _extract_docx_stream(stream: BinaryIO) -> str:
    try:
        import docx  # python-docx
    except ImportError as e:
        raise RuntimeError("Reading .docx requires python-docx (pip install python-docx).") from e
    document = docx.Document(stream)
    parts = [p.text for p in document.paragraphs]
    # Also pull text out of tables (common in resumes).
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    parts.append(cell.text)
    return "\n".join(parts)


# ---------------------------------------------------------------------------
# Structured profile extraction (lightweight heuristics, no ML)
# ---------------------------------------------------------------------------

_TITLE_KEYWORDS = [
    "software engineer", "software developer", "full stack developer", "full-stack developer",
    "frontend developer", "front-end developer", "backend developer", "back-end developer",
    "web developer", "mobile developer", "ios developer", "android developer",
    "data scientist", "data analyst", "data engineer", "machine learning engineer",
    "ml engineer", "ai engineer", "research scientist", "analytics engineer",
    "devops engineer", "site reliability engineer", "cloud engineer", "platform engineer",
    "security engineer", "qa engineer", "test engineer", "automation engineer",
    "product manager", "project manager", "program manager", "product owner",
    "business analyst", "systems analyst", "solutions architect", "software architect",
    "engineering manager", "technical lead", "team lead", "scrum master",
    "ux designer", "ui designer", "product designer", "graphic designer",
    "marketing manager", "digital marketer", "content strategist", "seo specialist",
    "sales manager", "account manager", "account executive", "customer success manager",
    "financial analyst", "accountant", "controller", "operations manager",
    # Research / academia / policy / other domains (not just tech):
    "ai ethics researcher", "ai researcher", "ai policy researcher", "research scientist",
    "research fellow", "research engineer", "postdoctoral researcher", "research associate",
    "research assistant", "researcher", "professor", "lecturer", "ethicist",
    "policy advisor", "policy analyst", "scientist", "economist", "writer", "editor",
    "journalist", "teacher", "nurse", "physician", "attorney", "lawyer", "paralegal",
    "recruiter", "human resources", "hr manager",
    # Generic fallbacks (kept last; word-boundary matched):
    "consultant", "developer", "engineer", "designer", "analyst", "manager",
    "administrator", "specialist", "coordinator", "director", "architect", "fellow",
]

_SENIORITY_PATTERNS = [
    ("lead", re.compile(r"\b(lead|principal|staff|head of|director|vp|chief)\b", re.I)),
    ("senior", re.compile(r"\b(senior|sr\.?|sr|expert)\b", re.I)),
    ("junior", re.compile(r"\b(junior|jr\.?|entry[- ]level|graduate|intern|trainee)\b", re.I)),
]


def _detect_titles(text: str) -> list[str]:
    lowered = text.lower()
    found: dict[str, int] = {}
    for title in _TITLE_KEYWORDS:
        # Word-boundary match so "engineer" doesn't match inside "engineering"
        # (e.g. "prompt engineering" must not be read as an Engineer role).
        m = re.search(rf"(?<![a-z]){re.escape(title)}(?![a-z])", lowered)
        if m and title not in found:
            found[title] = m.start()
    # Prefer more specific (longer) titles and earlier appearance.
    ordered = sorted(found.items(), key=lambda kv: (kv[1], -len(kv[0])))
    # De-duplicate generic titles that are substrings of a more specific one already found.
    result: list[str] = []
    for title, _ in ordered:
        if any(title != other and title in other for other in found):
            continue
        result.append(title)
    return result[:5]


def _detect_years_experience(text: str) -> int | None:
    # Look for explicit phrases like "8 years of experience".
    candidates = [int(m) for m in re.findall(r"(\d{1,2})\+?\s*years?\s+(?:of\s+)?experience", text, re.I)]
    if candidates:
        return max(candidates)
    # Otherwise estimate from a span of years mentioned (e.g. 2015 - 2024).
    years = [int(y) for y in re.findall(r"\b(19[89]\d|20[0-4]\d)\b", text)]
    if len(years) >= 2:
        span = max(years) - min(years)
        if 0 < span <= 50:
            return span
    return None


def _detect_seniority(text: str, years: int | None) -> str | None:
    head = text[:600]  # seniority cues are usually near the top (title/summary)
    for label, pattern in _SENIORITY_PATTERNS:
        if pattern.search(head):
            return label
    if years is not None:
        if years >= 8:
            return "senior"
        if years >= 3:
            return "mid"
        return "junior"
    return None


def _detect_location(text: str) -> str | None:
    # Heuristic: a "City, ST" or "City, Country" pattern in the first lines.
    for line in text.splitlines()[:15]:
        m = re.search(r"\b([A-Z][a-zA-Z.\- ]{2,30}),\s*([A-Z]{2}|[A-Z][a-zA-Z ]{2,30})\b", line)
        if m and "experience" not in line.lower():
            return f"{m.group(1).strip()}, {m.group(2).strip()}"
    return None


def _detect_name(text: str) -> str | None:
    """Best-effort: the candidate's name is usually the first real line of a CV."""
    for line in text.splitlines()[:6]:
        s = line.strip()
        if not s or "@" in s or any(ch.isdigit() for ch in s):
            continue
        # Drop a leading honorific for the word-count check, keep it in the result.
        core = re.sub(r"^(dr|mr|mrs|ms|prof)\.?\s+", "", s, flags=re.I)
        words = core.split()
        if 1 <= len(words) <= 4 and all(w[:1].isupper() for w in words if w):
            low = s.lower()
            if any(k in low for k in ("curriculum", "resume", "cv", "engineer", "developer",
                                      "manager", "analyst", "profile")):
                continue
            return s
    return None


def build_profile(text: str) -> CVProfile:
    """Build a structured profile from raw CV text."""
    text = text or ""
    skills = extract_skills(text)
    titles = _detect_titles(text)
    years = _detect_years_experience(text)
    seniority = _detect_seniority(text, years)
    location = _detect_location(text)

    # Build a sensible default search query: most specific title + top skills.
    query_bits: list[str] = []
    if titles:
        query_bits.append(titles[0])
    elif skills:
        query_bits.append(skills[0])
    suggested = " ".join(query_bits).strip()

    return CVProfile(
        raw_text=text,
        name=_detect_name(text),
        skills=skills,
        titles=titles,
        years_experience=years,
        location=location,
        seniority=seniority,
        suggested_keywords=suggested,
    )


def looks_empty(text: str) -> bool:
    """True if extracted text is too sparse to be a real CV.

    Common cause: a scanned / image-only PDF that has no embedded text layer.
    """
    if not text:
        return True
    alnum = sum(c.isalnum() for c in text)
    return alnum < 80


def parse_cv(path: str | Path) -> CVProfile:
    """Convenience: extract text from a file and build a structured profile."""
    return build_profile(extract_text(path))
